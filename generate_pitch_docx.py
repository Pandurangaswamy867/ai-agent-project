from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def create_pitch_docx():
    doc = Document()
    
    # Define Colors
    NAVY = RGBColor(0, 33, 71)
    ORANGE = RGBColor(255, 107, 0)

    # Title Slide Equivalent
    title = doc.add_heading('AI-Driven MSE Onboarding and Strategic Partner Mapping Ecosystem', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('BREAKING THE SILENCE: Voice-First AI Onboarding for India\'s 63M MSMEs')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = ORANGE
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('National Digital Public Good | ONDC Integration Platform').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. The Vision
    doc.add_heading('1. The Vision', level=1)
    doc.add_paragraph('To create a seamless, multilingual, and AI-verified gateway that empowers every Micro and Small Enterprise (MSE) in India to join the digital commerce revolution without language or technical barriers.')

    # 2. The Silent Crisis (Problem)
    doc.add_heading('2. The Silent Crisis: The Digital Divide', level=1)
    p = doc.add_paragraph('Digital exclusion in India isn’t a choice; it’s a language barrier. 63 Million MSMEs are the heart of Bharat, yet 90% remain digitally invisible due to:')
    doc.add_paragraph('Language Barriers: 90% of onboarding tools are English-centric.', style='List Bullet')
    doc.add_paragraph('Process Complexity: Manual forms and document verification are slow and error-prone.', style='List Bullet')
    doc.add_paragraph('Partner Isolation: MSEs lack intelligent discovery for ONDC fulfillment.', style='List Bullet')

    # 3. The AI Bridge (Solution)
    doc.add_heading('3. The AI Bridge', level=1)
    p = doc.add_paragraph('From Mother Tongue to Global Market in 60 Seconds.')
    doc.add_paragraph('Voice-First Onboarding: Speak naturally in 7+ Indian dialects to register.', style='List Bullet')
    doc.add_paragraph('Zero-Touch Compliance: Donut-based OCR extracts compliance data in real-time.', style='List Bullet')
    doc.add_paragraph('Semantic Matching: Intelligent algorithms find the perfect fulfillment partner based on capability, not just keywords.', style='List Bullet')

    # 4. Proprietary Tech Stack
    doc.add_heading('4. Proprietary Intelligence & Innovation', level=1)
    doc.add_paragraph('Heavyweight Brains. Featherweight Footprint.')
    doc.add_paragraph('Whisper-Tiny ASR: Optimized for regional Indian accents and low-fidelity audio.', style='List Bullet')
    doc.add_paragraph('BART Zero-Shot: Instant product categorization without expensive retraining.', style='List Bullet')
    doc.add_paragraph('Lazy-Loading Architecture: 65% reduction in server memory costs.', style='List Bullet')
    doc.add_paragraph('Native WAV Encoding: Crystal clear audio capture directly in the browser.', style='List Bullet')

    # 5. Market & Scale
    doc.add_heading('5. Market Opportunity', level=1)
    doc.add_paragraph('India\'s $100B+ Digital Commerce Opportunity.')
    doc.add_paragraph('Total Addressable Market: 63M+ MSEs.', style='List Bullet')
    doc.add_paragraph('Target Market Share: 20% of new ONDC integrations within 24 months.', style='List Bullet')
    doc.add_paragraph('Sector Agnostic: Ready for Agri, Textiles, Manufacturing, and beyond.', style='List Bullet')

    # 6. Business Model
    doc.add_heading('6. Business Model & Sustainability', level=1)
    doc.add_paragraph('Tiered SaaS for SNPs: Match-volume and management fees.', style='List Bullet')
    doc.add_paragraph('Success-Based Commissions: Percentage fee on fulfilled AI-matched orders.', style='List Bullet')
    doc.add_paragraph('Enterprise APIs: Licensing core AI modules to logistics and gov bodies.', style='List Bullet')
    doc.add_paragraph('Inclusive Access: Registration is always FREE for MSE owners.', style='List Bullet')

    # 7. Measurable Dominance (Benchmarks)
    doc.add_heading('7. Performance Benchmarks', level=1)
    doc.add_paragraph('94.5% OCR Precision on identity document extraction.', style='List Bullet')
    doc.add_paragraph('92% Voice Recognition accuracy in regional Indian dialects.', style='List Bullet')
    doc.add_paragraph('88% Matching Alignment with domain expert recommendations.', style='List Bullet')
    doc.add_paragraph('< 2.5s Response Time for end-to-end AI actions.', style='List Bullet')

    # 8. Governance & Security
    doc.add_heading('8. Trust, Security & Compliance', level=1)
    doc.add_paragraph('DPDP Act 2023: Fully compliant with Indian data privacy standards.', style='List Bullet')
    doc.add_paragraph('AES-256 Encryption: Banking-grade security for business documents.', style='List Bullet')
    doc.add_paragraph('Tamper-Proof Audit: Every AI intent logged for ministry oversight.', style='List Bullet')

    # 9. Roadmap
    doc.add_heading('9. The Roadmap to National Scale', level=1)
    doc.add_paragraph('Current: Successful pilots in Varanasi (Textiles) and Tamil Nadu (Handlooms).', style='List Bullet')
    doc.add_paragraph('Next: Expansion to 10+ Indic languages and AI-fraud detection.', style='List Bullet')
    doc.add_paragraph('Goal: Onboard 1 Million MSEs into the ONDC network by 2027.', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph('Contact: support@team-initiative.gov.in').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('AI_MSE_Pitch_Deck_Report.docx')
    print("Pitch Deck Word doc saved as AI_MSE_Pitch_Deck_Report.docx")

if __name__ == "__main__":
    create_pitch_docx()
