from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_expertise_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Core Expertise Areas & Technical Capabilities', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Section 1
    doc.add_heading('1. Multilingual AI & Natural Language Processing (NLP)', level=1)
    doc.add_paragraph('Speech-to-Text (ASR): Expertise in deploying OpenAI Whisper optimized for regional Indian dialects and accented speech.', style='List Bullet')
    doc.add_paragraph('Semantic Intelligence: Deep proficiency in using Sentence-Transformers and Zero-Shot (BART) models for intent-based navigation and partner matching.', style='List Bullet')
    doc.add_paragraph('Regional Entity Mapping: Advanced capability in normalizing Indic language transcripts into standardized industrial taxonomies (ONDC/MSME).', style='List Bullet')

    # Section 2
    doc.add_heading('2. Computer Vision & Document Intelligence', level=1)
    doc.add_paragraph('OCR-Free Document QA: Specialized knowledge in using Vision Transformers (Donut) for direct end-to-end data extraction from business certificates.', style='List Bullet')
    doc.add_paragraph('Automated Compliance: Expertise in building secure verification pipelines for Aadhaar, PAN, and Udyam certification with high precision.', style='List Bullet')

    # Section 3
    doc.add_heading('3. High-Performance AI Orchestration', level=1)
    doc.add_paragraph('Resource-Aware Deployment: Mastery of lazy-loading and model-on-demand architectures to minimize cloud infrastructure costs.', style='List Bullet')
    doc.add_paragraph('Hybrid Engine Design: Unique ability to blend probabilistic AI outputs with deterministic rule-based triggers for 100% reliability in critical workflows.', style='List Bullet')

    # Section 4
    doc.add_heading('4. Scalable Digital Public Good Engineering', level=1)
    doc.add_paragraph('API Ecosystems: Expert-level Python/FastAPI development for low-latency, ONDC-compliant backend services.', style='List Bullet')
    doc.add_paragraph('Inclusive UI/UX: Advanced frontend development (React/TS) focused on voice-driven accessibility and resilient offline data persistence.', style='List Bullet')
    doc.add_paragraph('Native Browser Integration: Proficiency in Web Audio API for real-time, client-side PCM audio encoding.', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph('This core expertise base ensures that the solution is not only technically advanced but also practically viable for national-scale deployment across Bharat.')

    doc.save('Core_Expertise_Areas.docx')
    print("Expertise document saved as Core_Expertise_Areas.docx")

if __name__ == "__main__":
    create_expertise_doc()
